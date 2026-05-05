import streamlit as st

# --- Fix text area visibility ---
st.markdown("""
<style>
textarea {
    color: white !important;
    background-color: #0D1B2A !important; /* deep navy background */
    font-family: 'DM Sans', sans-serif;
}
</style>
""", unsafe_allow_html=True)

# --- Custom CSS for branding ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@600&family=DM+Sans&display=swap');

    html, body, [class*="css"]  {
        font-family: 'DM Sans', sans-serif;
        color: #0D1B2A;
        background-color: #F9F7F4;
    }

    h1, h2, h3, h4 {
        font-family: 'Playfair Display', serif;
        color: #F4A922;
    }

    /* Button styling */
    .stButton>button {
        background-color: #4CAF82;
        color: #F9F7F4;
        border-radius: 8px;
        font-weight: bold;
    }

    .stButton>button:hover {
        background-color: #0D1B2A;
        color: #F4A922;
    }

    /* Tab styling */
    div[data-baseweb="tab"] {
        font-family: 'Playfair Display', serif;
        font-size: 16px;
        font-weight: 600;
        background-color: #0D1B2A;   /* deep navy */
        color: #F9F7F4;              /* off-white */
        border-radius: 6px 6px 0 0;
        margin-right: 4px;
        padding: 8px 16px;
    }

    div[data-baseweb="tab"]:hover {
        background-color: #4CAF82;   /* muted green */
        color: #F9F7F4;
    }

    div[data-baseweb="tab"][aria-selected="true"] {
        background-color: #F4A922;   /* warm gold */
        color: #0D1B2A;              /* deep navy text */
    }

    /* Privacy note styling */
    .stAlert {
        background-color: #4CAF82 !important;  /* muted green */
        color: #F4A922 !important;             /* warm gold text */
        border-radius: 8px;
        font-weight: bold;
    }

    /* Footer styling */
    .footer {
        text-align: center;
        margin-top: 40px;
        font-size: 14px;
        color: #0D1B2A;
    }
    .footer a {
        color: #4CAF82;
        text-decoration: none;
        font-weight: bold;
    }
    .footer a:hover {
        color: #F4A922;
    }

    /* Tagline styling */
    .tagline {
        position: absolute;
        top: 10px;
        right: 20px;
        font-family: 'Playfair Display', serif;
        font-size: 18px;
        font-weight: bold;
        color: #F4A922;
    }
    </style>
""", unsafe_allow_html=True)

# --- Toggle between Testing Mode and Live Mode ---
TESTING_MODE = True   # Change to False when you want to use OpenAI again

# --- Tagline ---
st.markdown('<div class="tagline">Stop applying. Start positioning.</div>', unsafe_allow_html=True)

# --- Tab Navigation ---
tab1, tab2, tab3, tab4, tab5 = st.tabs(["CV Review", "AI Coach", "Learning Hub", "Fun Corner", "CV Builder"])

# --- CV REVIEW TAB ---
with tab1:
    st.header("Get Your CV Reviewed")
    st.write("Upload OR paste your CV and get honest, specific feedback — like having a senior hiring manager read your CV over coffee.")

    # Privacy note
    st.info("🔒 Your CV content is used only to generate your review and is not stored on our servers. See our Privacy Policy.")

    uploaded_file = st.file_uploader("Upload CV File", type=["pdf", "docx"])
    cv_text = st.text_area("Paste CV Text")

    if uploaded_file is not None:
        try:
            cv_text = uploaded_file.read().decode("utf-8", errors="ignore")
        except Exception:
            st.error("⚠️ Could not read file. Please paste text instead.")

    if st.button("Review My CV"):
        if cv_text:
            if TESTING_MODE:
                # Dummy recruiter-style feedback (no API call)
                st.markdown(f"""
Thank you for sharing your CV, Candidate.

I'll be very direct, speaking as someone who's reviewed CVs for over 10 years:

- Your CV is visually outdated — heavy borders and shaded boxes make it look pre-2010.
- The layout is not ATS-readable; large sections may be ignored by recruiter systems.
- Font choices, spacing, and alignment are inconsistent and strain the eye.
- The opening headline is weak — it does not position you for a specific role.
- The synopsis is generic, repetitive, and filled with buzzwords without context.
- Core competencies are written as paragraphs instead of sharp, scannable skill statements.

In short: your experience may be strong, but the presentation is actively working against you.
                """)

                # Add "Review another CV" button
                if st.button("↺ Review another CV"):
                    st.experimental_rerun()
            else:
                st.error("⚠️ Live Mode is disabled until billing/API credits are available.")
        else:
            st.warning("⚠️ Please upload or paste your CV first.")

# --- AI COACH TAB ---
with tab2:
    st.header("CareerClariFi Coach")
    st.write("Ask about your career, CV, LinkedIn, job search strategy, or salary negotiation.")

    # Define common career questions + dummy answers
    questions = {
        "Why am I not getting shortlisted despite applying for months?":
            "Recruiters often filter by keywords and relevance. If your CV doesn’t mirror the job spec, ATS may skip it. Focus on tailoring your CV headline, summary, and skills to each role.",
        "How do I position myself for a career pivot?":
            "Highlight transferable skills. Frame your past achievements in terms of impact, not industry jargon. Show how your experience solves problems in the new field.",
        "Should I include a professional summary on my CV?":
            "Yes — but keep it sharp. 3–4 lines that state your role, years of experience, one quantifiable achievement, and the value you bring.",
        "How do I negotiate a higher salary offer?":
            "Anchor your request in market data. State your value clearly, then ask with confidence. Avoid vague phrases like 'I was hoping for more' — instead say 'Based on my experience and market benchmarks, I believe Rxxx is fair.'",
        "What should my LinkedIn headline say?":
            "Your headline should position you, not just list your job title. Example: 'Financial Planner | Helping SMEs build credibility & compliance | Xero Advisor'.",
        "Am I targeting the right companies and industries?":
            "Check alignment between your skills and the industry’s growth areas. Use LinkedIn job insights and industry reports to confirm demand before applying."
    }

    # Display clickable questions
    selected_question = st.radio("Choose a question:", list(questions.keys()))

    if st.button("Get Answer"):
        st.subheader("Answer")
        st.write(questions[selected_question])

    if st.button("↺ Reset AI Coach"):
        st.experimental_set_query_params()  # clears selection

# --- LEARNING HUB TAB ---
with tab3:
    st.header("Learning Hub")
    st.write("Free resources to sharpen your career game — from CV writing to salary negotiation.")

    # Search bar
    search_query = st.text_input("Search resources...", "")

    # Categories
    categories = ["All", "CV Writing", "LinkedIn", "Job Search Strategy", "Interview Prep", "Career Pivots", "Industry Insights"]
    active_category = st.radio("Filter by category:", categories, horizontal=True)

    # Dummy resources with full article text for testing
    resources = [
        {
            "title": "Why Your CV Keeps Getting Ignored (And How to Fix It)",
            "description": "The harsh truth about why most CVs get skipped — and the 5 structural changes that make recruiters stop scrolling.",
            "read_time": "6 min read",
            "category": "CV Writing",
            "content": """
### Why Your CV Keeps Getting Ignored (And How to Fix It)

Let's start with a hard truth: most CVs look the same. And in a stack of 200+ applications, "the same" means invisible.

#### The Reality of CV Screening
Recruiters in South Africa typically spend 6–10 seconds on an initial CV scan. That's not enough time to read your career history — it's enough time to feel whether your CV is worth reading.

**Here's what gets your CV ignored:**

1. **No Professional Summary (Or a Generic One)**  
   Fix: Write a summary that includes your years of experience, your specialisation, one quantifiable achievement, and the value you bring.

2. **Job Descriptions Instead of Achievements**  
   Fix: For each role, include 3–5 bullet points that show impact. Use the formula: Action + Context + Result.

3. **Poor Visual Hierarchy**  
   Fix: Use clear section headings, consistent fonts, plenty of white space, and a logical top-to-bottom flow.

4. **Missing Keywords**  
   Fix: Mirror key phrases from the job ad naturally in your CV — especially in your summary, skills, and experience sections.

5. **Too Long or Too Short**  
   Fix: The sweet spot for mid-to-senior professionals is 2–3 pages.

#### The Bottom Line
Your CV isn't just a document — it's your first negotiation with a potential employer. If it doesn't sell you in 10 seconds, everything else is irrelevant.

**Next step:** Run your CV through a fresh pair of eyes. Not your mom's eyes — a recruiter's eyes.
            """
        }
    ]

    # Track selected article in session_state
    if "selected_resource" not in st.session_state:
        st.session_state.selected_resource = None

    # If an article is selected, show full content
    if st.session_state.selected_resource:
        res = st.session_state.selected_resource
        st.subheader(res["title"])
        st.caption(f"📄 {res['category']} • ⏱️ {res['read_time']}")
        st.markdown(res["content"])

        # Monetization banner
        st.markdown("""
        <div style="background-color:#F9F7F4; text-align:center; padding:20px; margin-top:30px; border-radius:8px;">
            <h3 style="color:#0D1B2A;">Ready to put this into action?</h3>
            <p style="color:#0D1B2A;">Book a CV Revamp or LinkedIn session with JoyTee Holdings and get personalised, expert help.</p>
            <a href="https://wa.me/27600000000" target="_blank" style="background-color:#F4A922; color:#0D1B2A; padding:10px 20px; border-radius:6px; text-decoration:none; font-weight:bold;">
                → Book a ClariFi session from R350
            </a>
        </div>
        """, unsafe_allow_html=True)

        if st.button("↺ Back to Learning Hub"):
            st.session_state.selected_resource = None

    else:
        # Filter resources
        filtered_resources = []
        for r in resources:
            matches_category = active_category == "All" or r["category"] == active_category
            matches_search = not search_query or search_query.lower() in r["title"].lower() or search_query.lower() in r["description"].lower()
            if matches_category and matches_search:
                filtered_resources.append(r)

        # Display resource list
        if filtered_resources:
            for res in filtered_resources:
                st.subheader(res["title"])
                st.write(res["description"])
                st.caption(f"⏱️ {res['read_time']}")
                if st.button(f"Read: {res['title']}"):
                    st.session_state.selected_resource = res
        else:
            st.write("🔍 No resources found. Try a different search or category.")

        if st.button("↺ Reset Learning Hub"):
            st.session_state.selected_resource = None

# --- FUN CORNER TAB ---
with tab4:
    st.header("The Funny Corner 😂")
    st.write("Job searching is stressful. Take a break, have a laugh, and remember — you're not alone in this.")

    # Daily Memes (12+ jokes)
    st.subheader("📸 Daily Career Memes ZA")
    memes = [
        "‘They said it’s not you, it’s the budget’… then posted the same job two weeks later.",
        "Recruiter: ‘We’re looking for a unicorn.’ Me: ‘I’m just trying to pay rent.’",
        "Job spec: 3 years experience. Salary: 3 peanuts per month.",
        "Interview: ‘Where do you see yourself in 5 years?’ Me: ‘Hopefully not still waiting for feedback.’",
        "LinkedIn: ‘Open to work.’ Recruiter: ‘Open to exploitation.’",
        "Boss: ‘We’re a family here.’ Translation: ‘We don’t pay overtime.’",
        "Job ad: ‘Fast‑paced environment.’ Translation: ‘You’ll burn out in 3 months.’",
        "Recruiter ghosted me. I guess Casper works in HR now.",
        "‘We’ll keep your CV on file.’ That file is called the recycle bin.",
        "‘We want someone with 10 years experience in a tool released last year.’",
        "‘Culture fit’ = we want someone who laughs at the boss’s jokes.",
        "‘Exposure instead of pay.’ Sorry, exposure doesn’t cover rent.",
        "‘We’re looking for a rockstar.’ Translation: ‘We want you to work nights for free.’",
        "‘We’ll get back to you soon.’ Soon = never.",
        "‘Dynamic environment.’ Translation: ‘We have no processes.’",
        "‘We value work‑life balance.’ Translation: ‘We’ll email you at 11pm.’",
        "‘We’re agile.’ Translation: ‘We change our minds every week.’",
        "‘We’re hiring due to growth.’ Translation: ‘Everyone quit.’",
        "‘We’re looking for self‑starters.’ Translation: ‘We won’t train you.’",
        "‘We’re like a family.’ Translation: ‘We’ll guilt trip you into staying late.’"
    ]
    meme_index = st.number_input("Swipe memes (1–20):", min_value=1, max_value=len(memes), value=1)
    st.write(memes[meme_index - 1])

    # Office Trivia / Quiz (6 hilarious questions)
    st.subheader("🎲 Office Trivia")
    trivia_questions = {
        "What does ‘fast‑paced environment’ really mean?": [
            "You’ll burn out in 3 months.",
            "You’ll be running to HR daily.",
            "It’s code for chaos."
        ],
        "What does ‘culture fit’ usually mean?": [
            "Laugh at the boss’s jokes.",
            "Pretend to enjoy team‑building karaoke.",
            "Agree that free coffee is a benefit."
        ],
        "What does ‘competitive salary’ mean?": [
            "They compete to see who gets paid less.",
            "It’s a mystery box.",
            "It’s below market, but they’ll smile."
        ],
        "What does ‘great exposure’ mean?": [
            "You’ll be exposed to poverty.",
            "You’ll be exposed to burnout.",
            "You’ll be exposed to unpaid overtime."
        ],
        "What does ‘dynamic environment’ mean?": [
            "Processes change every week.",
            "Your boss changes their mind hourly.",
            "You’ll be dynamically stressed."
        ],
        "What does ‘family culture’ mean?": [
            "They’ll guilt trip you into staying late.",
            "They’ll forget your birthday but remember deadlines.",
            "They’ll call you ‘family’ until retrenchment."
        ]
    }

    selected_question = st.selectbox("Pick a trivia question:", list(trivia_questions.keys()))
    selected_answer = st.radio("Choose your answer:", trivia_questions[selected_question])
    if st.button("Reveal Truth"):
        st.success(f"😂 Correct! {selected_answer}")

    # Footer note inside Funny Corner
    st.markdown("""
    <div style="text-align:center; margin-top:20px; font-size:12px; color:#0D1B2A;">
        Made with ❤️ and a bit of eish energy. Powered by JoyTee Holdings.
    </div>
    """, unsafe_allow_html=True)

    if st.button("↺ Reset Fun Corner"):
        st.session_state.clear()

import requests
from docx import Document
import io
import streamlit as st

# Function to call Hugging Face API safely
def generate_revamp(cv_text, job_desc):
    prompt = f"""
    Revamp the CV below to align with the job description:
    - New 45-50 word professional summary
    - Extract 6 core skills, 4 technical skills, 2 soft skills
    - Rewrite experience bullets (9 words each, aligned to JD)
    - Keep education/certifications intact
    CV:
    {cv_text}
    Job Description:
    {job_desc}
    """

    headers = {"Authorization": f"Bearer {st.secrets['huggingface']['api_key']}"}
    payload = {"inputs": prompt}

    response = requests.post(
        "https://api-inference.huggingface.co/models/google/flan-t5-large",  # ✅ supported model
        headers=headers, json=payload
    )

    try:
        data = response.json()
        if isinstance(data, list) and "generated_text" in data[0]:
            return data[0]["generated_text"]
        elif isinstance(data, dict) and "generated_text" in data:
            return data["generated_text"]
        else:
            return f"⚠️ Unexpected response format: {data}"
    except Exception as e:
        return f"⚠️ Error decoding response: {str(e)}\nRaw response: {response.text}"

# --- CV BUILDER TAB ---
with tab5:
    st.header("ATS CV Builder")
    st.write("Upload or paste your CV, paste the job description, and build a ready-to-apply ATS CV.")

    uploaded_cv = st.file_uploader("Upload CV File (PDF/Word)", type=["pdf", "docx"])
    pasted_cv = st.text_area("Or paste your CV text here")
    job_desc = st.text_area("Paste job description here")

    template_choice = st.selectbox("Choose a template:", 
                                   ["Classic ATS", "Modern Professional", "Strategic Pivot"])

    if st.button("Generate ATS CV"):
        if pasted_cv or uploaded_cv:
            st.subheader(f"Revamped CV: {template_choice} Template")

            # Get CV text
            cv_text = pasted_cv if pasted_cv else uploaded_cv.read().decode("utf-8", errors="ignore")

            # Call Hugging Face to generate revamped CV
            revamped_output = generate_revamp(cv_text, job_desc)

            # Show editable text area with AI output
            st.text_area("Revamped CV", revamped_output, height=400)

            # --- Create Word file ---
            doc = Document()
            doc.add_heading('ATS Compliant CV', 0)
            doc.add_paragraph(revamped_output)

            word_buffer = io.BytesIO()
            doc.save(word_buffer)
            word_buffer.seek(0)

            st.download_button(
                label="⬇️ Download CV (Word)",
                data=word_buffer,
                file_name="ATS_CV.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # PDF export disabled until fpdf2 is installed
            st.info("PDF export will be enabled once fpdf2 is installed. For now, use Word download.")
        else:
            st.error("Please upload or paste your CV first.")
            
            # Premium upsell
            st.markdown("""
            <div style="background-color:#F9F7F4; text-align:center; padding:20px; margin-top:30px; border-radius:8px;">
                <h3 style="color:#0D1B2A;">Want industry-specific positioning?</h3>
                <p style="color:#0D1B2A;">Our premium Career Branding & Positioning service goes beyond ATS — we research the right fit for your next progression.</p>
                <a href="https://wa.me/27600000000" target="_blank" style="background-color:#F4A922; color:#0D1B2A; padding:10px 20px; border-radius:6px; text-decoration:none; font-weight:bold;">
                    → Book Career Branding & Positioning
                </a>
            </div>
            """, unsafe_allow_html=True)
